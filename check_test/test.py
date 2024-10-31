from bs4 import BeautifulSoup
from docx import Document
import os
import win32com.client as win32


# 巡检内容列表
inspection_items = [
    "异常特权账户",
    "异常远程账户",
    "空密码账户",
    "爆破记录",
    "异常计划任务",
    "CPU使用率 < 80%",
    "内存使用率 < 80%",
    "TOP10进程信息",
    "僵尸进程",
    "服务器时间",
    "防火墙状态",
    "ESTABLISHED < 1000",
    "磁盘分区占用 < 80%",
]

def checked_all_correct(row):
    row.cells[3].paragraphs[0].clear().add_run('✔')
    row.cells[4].paragraphs[0].clear()
    row.cells[5].paragraphs[0].clear()

def checked_wrong_with_msg(row,msg):
    row.cells[3].paragraphs[0].clear().add_run('✘')
    row.cells[4].paragraphs[0].add_run(msg)


# 定义各个检查函数
def check_abnormal_privileged_accounts(row,check_items):
    item = check_items['用户安全审计'][0]
    # print(item)
    other = str(item).replace("特权用户列表：",'')

    if '' == other:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,other)


def check_abnormal_remote_accounts(row,check_items):
    item = check_items['用户安全审计'][1]
    # print(item)
    other = str(item).replace("可以远程登陆的用户列表：",'').replace("root",'')

    if '' == other:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,other)


def check_empty_password_accounts(row,check_items):
    item = check_items['用户安全审计'][2]
    # print(item)
    other = str(item).replace("密码为空的用户列表：",'')

    if '' == other:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,other)


def check_brute_force_records(row,check_items):
    item = check_items['登陆失败记录']
    # print(item)
    # print(len(item))

    if 6 == len(item):
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,"请手动检查")


def check_abnormal_scheduled_tasks(row,check_items):
    item = check_items['系统安全审计']
    # print(item)
    index = list(item).index("当前用户计划任务列表：")
    # print(start)
    task_list = []
    for i in range(index,len(item)):
        task_list.append(item[i])
    # print(task_list)

    if row.cells[3].paragraphs[0].text == '✘':
        row.cells[3].paragraphs[0].clear()
        row.cells[4].paragraphs[0].add_run("请手动检查")
    else:
        return


def check_cpu_usage(row,check_items):
    item = check_items['系统资源巡检区'][0]
    # print(item)
    max_usage = 80
    usage = item.replace('CPU使用率：','').replace('%','')
    cpu_usage = float(usage)
    # print(cpu_usage)

    if max_usage >= cpu_usage:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,item)


def check_memory_usage(row,check_items):
    item = check_items['配置信息'][8] # html可能修改位置 可参考check_zombie_processes check_established_connections
    # print(item)
    max_usage = 80
    usage = item.replace('内存使用率：','').replace('%','')
    ram_usage = float(usage)
    # print(ram_usage)

    if max_usage >= ram_usage:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,item)


def check_top10_processes(row,check_items):
    return


def check_zombie_processes(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统当前僵尸进程数' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = int(item.replace("系统当前僵尸进程数：",''))
    # print(number)

    if 0 == number:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,item)


def check_server_time(row,check_items):
    return


def check_firewall_status(row,check_items):
    return


def check_established_connections(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统 established socket数量' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = int(item.replace("系统 established socket数量: ",''))
    # print(number)
    max_connections = 1000

    if max_connections > number:
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,item)


def check_disk_partition_usage(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统磁盘分区存储使用情况' in s), None)
    # print(index)
    end = next((i for i, s in enumerate(item) if '系统当前进程数' in s), None)
    infos = []
    for i in range(index,end):
        infos.append(item[i])
    # print(infos)

    disk_list = []
    # 只保留包含 "%" 和 "/" 的行
    filtered_infos = [info for info in infos if '%' in info and '/' in info]
    for info in filtered_infos:
        parts = info.split()  # 将每一行按照空格分割
        if len(parts) > 5:  # 确保这一行有足够的字段
            try:
                if "%" in parts[5]:
                    use_percentage = parts[5].strip('%')  # 去掉 '%'
                elif "%" in parts[4]:
                    use_percentage = parts[4].strip('%')  # 去掉 '%'
                disk_list.append(float(use_percentage))  # 转换为整数
            except ValueError:
                print(disk_list)
                print(use_percentage)

    # print(disk_list)
    max_percentage = 80

    if all(max_percentage > value for value in disk_list):
        checked_all_correct(row)
    else:
        checked_wrong_with_msg(row,"请手动检查")


# 加载 HTML 文件
def load_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# 解析 HTML 获取检查项和服务器 IP
def parse_html_for_check_items(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    check_items = {}
    
    # 遍历所有 <h2> 标签
    for h2 in soup.find_all('h2'):
        section_title = h2.get_text(strip=True)
        
        # 找到紧跟着的 <pre> 标签
        pre = h2.find_next('pre')
        if pre:
            pre_content = pre.get_text(strip=True)
            if pre_content:
                # 根据换行符分割内容
                lines = pre_content.split('\n')
                # 将每个条目添加到检查项字典中，标题为键
                check_items[section_title] = [line.strip() for line in lines if line.strip()]
    
    # print(check_items)
    return check_items

def insert_file_into_table_cell(doc_path, file_path, ip_address):
    # 获取当前脚本所在的文件夹路径
    current_directory = os.getcwd()

    # 组合文件的完整路径
    doc_path = os.path.join(current_directory, doc_path)
    file_path = os.path.join(current_directory, file_path)
    # 打开 Word 应用
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False  # 设置为True可以看到Word界面

    try:
        # 打开现有的 Word 文档
        doc = word_app.Documents.Open(doc_path)

        # 遍历文档中的所有表格
        for table in doc.Tables:
             # 遍历表格中的每一个单元格
            for cell in table.Range.Cells:
                # 获取当前单元格所在的行号和列号
                row_index = cell.RowIndex
                col_index = cell.ColumnIndex

                # 如果是第一列，检查该单元格是否包含指定的 IP 地址
                if col_index == 1 and ip_address in cell.Range.Text.strip():
                    # print(f"与 IP 地址 {ip_address} 匹配的表格行: 第 {row_index} 行")

                    # 获取对应行的右侧单元格（第二列）
                    right_cell = table.Cell(Row=row_index, Column=2)

                    # 在右侧单元格中插入文件（HTML 或其他文件）
                    right_cell.Range.InlineShapes.AddOLEObject(
                        ClassType="",
                        FileName=file_path,  # 文件路径
                        LinkToFile=False,
                        DisplayAsIcon=True,
                        IconLabel=os.path.basename(file_path)  # 文件图标下方显示的文件名
                    )
                    # print(f"文件已成功插入到与 IP 地址 {ip_address} 对应的右侧单元格中。")

        # 保存并关闭文档
        doc.Save()
    except Exception as e:
        print(f"打开或处理文档时出错: {e}")

    finally:
        # 确保关闭文档和 Word 应用
        if 'doc' in locals():  # 确保 doc 已成功打开
            doc.Close()
        word_app.Quit()

def check_matching_rows(matching_rows):
    for row in matching_rows:
            # print([cell.text for cell in row.cells])
            check_row = row.cells[2].text.strip()  # 巡检内容列
            # print(check_row)
            
            if check_row in inspection_items:
                
                if check_row == inspection_items[0]:
                    check_abnormal_privileged_accounts(row,check_items)
                    
                elif check_row == inspection_items[1]:
                    check_abnormal_remote_accounts(row,check_items)

                elif check_row == inspection_items[2]:
                    check_empty_password_accounts(row,check_items)
                    
                elif check_row == inspection_items[3]:
                    check_brute_force_records(row,check_items)

                elif check_row == inspection_items[4]:
                    check_abnormal_scheduled_tasks(row,check_items)

                elif check_row == inspection_items[5]:
                    check_cpu_usage(row,check_items)

                elif check_row == inspection_items[6]:
                    check_memory_usage(row,check_items)

                elif check_row == inspection_items[7]:
                    check_top10_processes(row,check_items)

                elif check_row == inspection_items[8]:
                    check_zombie_processes(row,check_items)

                elif check_row == inspection_items[9]:
                    check_server_time(row,check_items)

                elif check_row == inspection_items[10]:
                    check_firewall_status(row,check_items)

                elif check_row == inspection_items[11]:
                    check_established_connections(row,check_items)

                elif check_row == inspection_items[12]:
                    check_disk_partition_usage(row,check_items)


# DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
def matching_docx_tables(new_path, ip_address,check_items):
    doc = Document(new_path)
    matching_rows = []
    for table in doc.tables:
        for row in table.rows:
            # 只在表格的前两列中查找匹配的 IP 地址
            if len(row.cells) > 2 and ip_address in row.cells[1].text:
                matching_rows.append(row)
    # 匹配的行内容
    if len(row.cells) == 2 and row.cells[1].text == '':
        row.cells[1].paragraphs[0].clear()
    
    if matching_rows:
        check_matching_rows(matching_rows)

    # 保存修改后的文档
    try:
        doc.save(new_path)
    except PermissionError:
        print("保存文档时出错，请确保 Word 文档未被占用，然后重新运行脚本。")
        exit(1)



if __name__ == "__main__":
    # 文件夹路径和 DOCX 文件路径
    html_folder_path = './test_file'
    docx_file_path = '浙江中医药大学运维巡检-20240712.docx'
    
    # 提示用户手动关闭 Word 程序
    input(f"请确保已关闭 Word 文档，然后按 Enter 键继续...")

    doc = Document(docx_file_path)
    new_path = 'checked_' + docx_file_path
    doc.save(new_path)
    # 遍历文件夹中的所有 HTML 文件
    for html_filename in os.listdir(html_folder_path):
        if html_filename.endswith('.html'):
            html_file_path = os.path.join(html_folder_path, html_filename)
            
            # 提取 IP 地址（假设文件名以 IP 开头）
            ip_address = html_filename.split('_')[0]
            
            # 加载和解析 HTML 文件
            html_content = load_html(html_file_path)
            check_items = parse_html_for_check_items(html_content)
            
            # 打印 HTML 文件中的检查项
            # print(f"检查项 (来自 {html_filename}):")
            # for item in check_items:
            #     print(item)
            # print("\n")
            
            # 更新 DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
            if ip_address:
                print(f"与 IP 地址 {ip_address} 匹配的表格内容:")
                # matching_rows = 
                
                matching_docx_tables(new_path, ip_address,check_items)
                # insert_file_into_table_cell(new_path,html_file_path,ip_address)
                # print(matching_rows)
                # rows_check(matching_rows)
                print("\n")
    # 打开修改后的文档
    os.startfile(new_path)