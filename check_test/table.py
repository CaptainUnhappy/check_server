import os
from docx import Document

# 文档路径
doc_path = 'table.docx'
modified_doc_path = 'modified_file.docx'

# 提示用户手动关闭 Word 程序
input(f"请确保已关闭{doc_path}与{modified_doc_path}，然后按 Enter 键继续...")

# 加载文档
try:
    doc = Document(doc_path)
except PermissionError:
    print(f"无法访问文档，请确保{doc_path}与{modified_doc_path}未被占用，然后重新运行脚本。")
    exit(1)

# 巡检内容列表
inspection_items = [
    "异常特权账户",
    "异常远程账户",
    "空密码账户",
    "爆破记录",
    "异常计划任务",
    "CPU使用率 < 80%",
    "内存使用率 < 80%",
    "TOP10进程信息",
    "僵尸进程",
    "服务器时间",
    "防火墙状态",
    "ESTABLISHED < 1000",
    "磁盘分区占用 < 80%",
]

# 遍历文档中的所有表格
for table in doc.tables:
    # 遍历表格中的行
    for row in table.rows:
        # 打印当前行
        print([cell.text for cell in row.cells])
        
        # 检查行中是否有足够的单元格
        if len(row.cells) > 2:
            # 提取用于条件检查的单元格值
            server_ip = row.cells[1].text.strip()  # 服务器 IP 列
            check_item = row.cells[2].text.strip()  # 巡检内容列

            for i in inspection_items:
                # 检查特定条件
                if i == check_item:
                    # 在相应的单元格中设置勾选标记，并保留现有格式
                    row.cells[3].paragraphs[0].clear().add_run('✔')
                    row.cells[4].paragraphs[0].add_run('4')
                    row.cells[5].paragraphs[0].add_run('5')

# 保存修改后的文档
try:
    doc.save(modified_doc_path)
except PermissionError:
    print("保存文档时出错，请确保 Word 文档未被占用，然后重新运行脚本。")
    exit(1)

# 打开修改后的文档
os.startfile(modified_doc_path)
