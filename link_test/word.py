import argparse
import re
import shutil
from bs4 import BeautifulSoup
from docx import Document
import os
import win32com.client as win32
import logging
from datetime import datetime

# 获取当前时间并格式化为文件夹名
current_time = datetime.now().strftime("%Y%m%d_%H%M%S")

# 获取当前工作目录，并创建一个以时间命名的文件夹
current_dir = os.getcwd()
save_folder = os.path.join(current_dir, 'output', current_time)
os.makedirs(save_folder, exist_ok=True)

# 设置日志文件的路径
log_file = os.path.join(save_folder, f"{current_time}.log")

# 配置日志记录，确保日志输出到文件和控制台，且文件以UTF-8编码
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# 创建日志文件处理器，保存为UTF-8编码
file_handler = logging.FileHandler(log_file, encoding='utf-8')
file_handler.setLevel(logging.INFO)

# 创建控制台处理器
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)

# 设置日志格式
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# 将处理器添加到logger
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# 巡检内容列表
inspection_items = [
    "异常特权账户",
    "异常远程账户",
    "空密码账户",
    "爆破记录",
    "异常计划任务",
    "CPU使用率 < 80%",
    "内存使用率 < 80%",
    "TOP10进程信息",
    "僵尸进程",
    "服务器时间",
    "防火墙状态",
    "ESTABLISHED < 1000",
    "磁盘分区占用 < 80%",
]

def checked_all_correct(row,msg):
    row.cells[3].paragraphs[0].clear().add_run('✔')
    row.cells[4].paragraphs[0].clear()
    row.cells[5].paragraphs[0].clear()
    # 可注释不打印
    logger.info(f'    √ {msg}') # cmd中显示勾

def checked_wrong_with_msg(row,msg):
    row.cells[3].paragraphs[0].clear().add_run('✘')
    row.cells[4].paragraphs[0].add_run(msg)
    logger.warning(f'× {msg}') # cmd中显示叉


# 定义各个检查函数
def check_abnormal_privileged_accounts(row,check_items):
    item = check_items['用户安全审计'][0]
    # print(item)
    other = str(item).replace("特权用户列表：",'')

    msg = '异常特权账户 '
    if '' == other:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+other)


def check_abnormal_remote_accounts(row, check_items):
    item = check_items['用户安全审计'][2]
    other = str(item).replace("密码为空的用户列表：", '')

    msg = '异常远程账户 '
    if '' == other:
        checked_all_correct(row, msg)
    else:
        checked_wrong_with_msg(row, msg + other)


def check_empty_password_accounts(row,check_items):
    item = check_items['用户安全审计'][2]
    # print(item)
    other = str(item).replace("密码为空的用户列表：",'')

    msg = '空密码账户 '
    if '' == other:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+other)


def check_brute_force_records(row, check_items):
    item = check_items['登陆失败记录']
    # print("原始记录:", item)

    cleaned_list = []
    recording = False
    for line in item:
        # 遇到 'IP              Failes' 开始记录
        if line == 'IP              Failes':
            recording = True
            continue
        # 遇到 '爆破主机root账号的可疑IP记录:' 停止记录第一部分
        if line == '爆破主机root账号的可疑IP记录:':
            recording = False
            continue
        # 如果正在记录，将内容添加到清理后的列表中
        if recording:
            cleaned_list.append(line)

    msg = '爆破记录 '
    if len(cleaned_list) == 0:
        checked_all_correct(row, msg)
    else:
        checked_wrong_with_msg(row, msg + "请手动检查")


def check_abnormal_scheduled_tasks(row,check_items):
    item = check_items['系统安全审计']
    # print(item)
    index = list(item).index("当前用户计划任务列表：")
    # print(start)
    task_list = []
    for i in range(index,len(item)):
        task_list.append(item[i])
    # print(task_list)

    msg = '异常计划任务 '
    if row.cells[3].paragraphs[0].text == '✔':
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+"请手动检查")


def check_cpu_usage(row,check_items):
    item = check_items['系统资源巡检区'][0]
    # print(item)
    max_usage = 80
    usage = item.replace('CPU使用率：','').replace('%','')
    if usage == '':
        usage = 0
    cpu_usage = float(usage)
    # print(cpu_usage)

    msg = 'CPU使用率 < 80% '
    if max_usage >= cpu_usage:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_memory_usage(row,check_items):
    item = check_items['配置信息']
    # print(item)
    index = next((i for i, s in enumerate(item) if '内存使用率' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    max_usage = 80
    usage = item.replace('内存使用率：','').replace('%','')
    ram_usage = float(usage)
    # print(ram_usage)

    msg = '内存使用率 < 80% '
    if max_usage >= ram_usage:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_top10_processes(row,check_items):
    # 可注释不打印
    logger.info(' 请手动检查 TOP10进程信息')
    return


def check_zombie_processes(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统当前僵尸进程数' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = item.replace("系统当前僵尸进程数：",'')
    # print(number)
    if number == '':
        number = 0
    number = int(number)
    msg = '僵尸进程 '
    if 0 == number:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_server_time(row,check_items):
    checked_all_correct(row,'服务器时间')


def check_firewall_status(row,check_items):
    checked_all_correct(row,'防火墙状态')


def check_established_connections(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统 established socket数量' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = int(item.replace("系统 established socket数量: ",''))
    # print(number)
    max_connections = 1000
    msg = 'ESTABLISHED < 1000 '
    if max_connections > number:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_disk_partition_usage(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统磁盘分区存储使用情况' in s), None)
    # print(index)
    end = next((i for i, s in enumerate(item) if '系统当前进程数' in s), None)
    infos = []
    for i in range(index,end):
        infos.append(item[i])
    # print(infos)

    disk_list = []
    # 只保留包含 "%" 和 "/" 的行
    filtered_infos = [info for info in infos if '%' in info and '/' in info]
    for info in filtered_infos:
        parts = info.split()  # 将每一行按照空格分割
        if len(parts) > 5:  # 确保这一行有足够的字段
            try:
                if "%" in parts[5]:
                    use_percentage = parts[5].strip('%')  # 去掉 '%'
                elif "%" in parts[4]:
                    use_percentage = parts[4].strip('%')  # 去掉 '%'
                disk_list.append(float(use_percentage))  # 转换为整数
            except ValueError:
                logger.error(f'取值错误: {disk_list} \n位于该行:{infos}')

    # print(disk_list)
    max_percentage = 80
    msg = '磁盘分区占用 < 80% '
    if all(max_percentage > value for value in disk_list):
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+"请手动检查")


# 加载 HTML 文件
def load_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# 解析 HTML 获取检查项和服务器 IP
def parse_html_for_check_items(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    check_items = {}
    
    # 遍历所有 <h2> 标签
    for h2 in soup.find_all('h2'):
        section_title = h2.get_text(strip=True)
        
        # 找到紧跟着的 <pre> 标签
        pre = h2.find_next('pre')
        if pre:
            pre_content = pre.get_text(strip=True)
            if pre_content:
                # 根据换行符分割内容
                lines = pre_content.split('\n')
                # 将每个条目添加到检查项字典中，标题为键
                check_items[section_title] = [line.strip() for line in lines if line.strip()]
    
    # print(check_items)
    return check_items

def check_matching_rows(matching_rows):
    for row in matching_rows:
            # print([cell.text for cell in row.cells])
            check_row = row.cells[2].text.strip()  # 巡检内容列
            # print(check_row)
            
            if check_row in inspection_items:
                
                if check_row == inspection_items[0]:
                    check_abnormal_privileged_accounts(row,check_items)
                    
                elif check_row == inspection_items[1]:
                    check_abnormal_remote_accounts(row,check_items)

                elif check_row == inspection_items[2]:
                    check_empty_password_accounts(row,check_items)
                    
                elif check_row == inspection_items[3]:
                    check_brute_force_records(row,check_items)

                elif check_row == inspection_items[4]:
                    check_abnormal_scheduled_tasks(row,check_items)

                elif check_row == inspection_items[5]:
                    check_cpu_usage(row,check_items)

                elif check_row == inspection_items[6]:
                    check_memory_usage(row,check_items)

                elif check_row == inspection_items[7]:
                    check_top10_processes(row,check_items)

                elif check_row == inspection_items[8]:
                    check_zombie_processes(row,check_items)

                elif check_row == inspection_items[9]:
                    check_server_time(row,check_items)

                elif check_row == inspection_items[10]:
                    check_firewall_status(row,check_items)

                elif check_row == inspection_items[11]:
                    check_established_connections(row,check_items)

                elif check_row == inspection_items[12]:
                    check_disk_partition_usage(row,check_items)


# DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
def matching_docx_tables(new_path, ip_address,check_items):
    doc = Document(new_path)
    matching_rows = []
    for table in doc.tables:
        for row in table.rows:
            # 只在表格的前两列中查找匹配的 IP 地址
            if len(row.cells) > 2 and ip_address in row.cells[1].text:
                matching_rows.append(row)
    # 匹配的行内容
    # print(matching_rows)
    
    if matching_rows:
        check_matching_rows(matching_rows)
    else:
        logger.error('!!!未找到匹配{ip_address}的表格!!!')

    # 保存修改后的文档
    try:
        doc.save(new_path)
    except PermissionError:                                                                                          
        logger.error("保存文档时出错，请确保 Word 文档未被占用，然后重新运行")
        exit(1)

def update_docx_info(doc_path, html_file_path, ip_address):
    # 获取当前脚本所在的文件夹路径
    current_directory = os.getcwd()

    # 组合文件的完整路径
    doc_path = os.path.join(current_directory, doc_path)
    html_file_path = os.path.join(current_directory, html_file_path)
    # 获取用户信息
    user_info = load_user_info()

    # 打开 Word 应用
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False  # 设置为True可以看到Word界面

    try:
        # 打开现有的 Word 文档
        doc = word_app.Documents.Open(doc_path)

        global first_update
        if first_update:
            # 更新大标题
            if doc.Shapes.Count > 0:
                first_shape = doc.Shapes[0]
                if hasattr(first_shape, "TextFrame") and first_shape.TextFrame.HasText:
                    original_text = first_shape.TextFrame.TextRange.Text.strip()  # 原始文本
                    try:
                        quarter = user_info['quarter']
                        if quarter == '1':
                            quarter_str = '一'
                        elif quarter == '2':
                            quarter_str = '二'
                        elif quarter == '3':
                            quarter_str = '三'
                        elif quarter == '4':
                            quarter_str = '四'
                        # 使用正则表达式替换符合模式的文本
                        new_text = re.sub(r"\d{4}年第.*季度运维巡检", f"{datetime.now().year}年第{quarter_str}季度运维巡检", original_text)
                        if new_text != original_text:
                            first_shape.TextFrame.TextRange.Text = new_text  # 更新文本
                            logger.info(f"已更新 大标题为 {new_text}")
                    except Exception as e:
                        logger.error(f"更新大标题时出错: {e}")                

            # 更新第一个表格的内容
            if doc.Tables.Count > 0:  # 确保文档至少有一个表格
                first_table = doc.Tables[0]  # 获取第一个表格
                row_index = 2  # 第二行
                
                try:
                    # 更新日期
                    date_str = datetime.now().strftime('%Y/%m/%d')
                    first_table.Cell(Row=row_index, Column=1).Range.Text = date_str
                    # 更新季度
                    quarter_str = 'Q'+user_info['quarter']
                    first_table.Cell(Row=row_index, Column=2).Range.Text = quarter_str
                    # 更新姓名
                    first_table.Cell(Row=row_index, Column=5).Range.Text = user_info['name']
                    # 更新电话
                    first_table.Cell(Row=row_index, Column=6).Range.Text = user_info['number']
                    
                    logger.info(f"已更新 日期={date_str}, 季度={quarter_str}, 姓名={user_info['name']}, 电话={user_info['number']}")
                except Exception as e:
                    logger.error(f"更新信息时出错: {e}")
            first_update = False


        # 遍历文档中的所有表格
        for table in doc.Tables:
            if "IP实例" in table.Cell(1, 1).Range.Text and "html报告" in table.Cell(1, 2).Range.Text:
                for row in range(2, len(table.Rows) + 1):  # 从第二行开始遍历，因为第一行是表头
                    # 获取左侧单元格的 IP 地址
                    ip_cell = table.Cell(row, 1).Range.Text.strip()

                    # 如果 IP 地址匹配
                    if ip_address in ip_cell:
                        # 获取对应行的右侧单元格(第二列)
                        right_cell = table.Cell(Row=row, Column=2)

                        # 删除原有的嵌入对象
                        for shape in right_cell.Range.InlineShapes:
                            shape.Delete()  # 删除该单元格中的所有嵌入对象

                        # 检查HTML文件是否存在
                        if os.path.exists(html_file_path):
                            # 在右侧单元格中插入HTML文件
                            right_cell.Range.InlineShapes.AddOLEObject(
                                ClassType="htmlfile",
                                FileName=html_file_path,  # HTML文件路径
                                LinkToFile=False,
                                DisplayAsIcon=True,
                                # IconFileName="C:\\Windows\\System32\\shell32.dll",
                                IconFileName="C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe",
                                IconIndex=0,
                                IconLabel=os.path.basename(html_file_path)  # 文件图标下方显示的文件名
                            )
                            # 可注释不打印
                            logger.info(f"文件已成功插入到与 IP 地址 {ip_address} 对应的第 {row} 行右侧单元格中。")
                            return
                        else:
                            logger.error(f"未找到IP地址 {ip_address} 对应的HTML文件: {html_file_path}")

        # 保存并关闭文档
        doc.Save()
    except Exception as e:
        logger.error(f"打开或处理文档时出错: {e}")

    finally:
        # 确保关闭文档和 Word 应用
        if 'doc' in locals():  # 确保 doc 已成功打开
            doc.Close()
        word_app.Quit()
        
# 复制 DOCX 和 HTML 文件到 save_folder
def copy_files_to_save_folder(docx_file_path, html_folder_path, save_folder):
    # 确保 save_folder 存在
    os.makedirs(save_folder, exist_ok=True)
    
    # 复制 DOCX 文件
    docx_dest_path = os.path.join(save_folder, os.path.basename(docx_file_path))
    shutil.copy(docx_file_path, docx_dest_path)  # 将 DOCX 文件复制到目标路径
    
    # 复制 DOCX 文件(用于检查)
    check_docx_dest_path = os.path.join(save_folder, 'check_' + os.path.basename(docx_file_path))
    shutil.copy(docx_file_path, check_docx_dest_path)  # 创建检查用的 DOCX 文件副本
    logger.info(f"已备份并新建 DOCX 文件到 {check_docx_dest_path}")
    
    # 创建保存 HTML 文件的文件夹
    html_dest_folder = os.path.join(save_folder, 'html_files')
    os.makedirs(html_dest_folder, exist_ok=True)  # 如果文件夹不存在则创建

    # 复制 HTML 文件
    for html_filename in os.listdir(html_folder_path):
        if html_filename.endswith('.html'):
            html_src_path = os.path.join(html_folder_path, html_filename)
            html_dest_path = os.path.join(html_dest_folder, html_filename)
            shutil.copy(html_src_path, html_dest_path)  # 将 HTML 文件复制到目标路径
            logger.info(f"已复制 HTML 文件到 {html_dest_path}")
    
    # 返回检查用 DOCX 文件路径和 HTML 文件夹路径
    return check_docx_dest_path, html_dest_folder

# 添加用户信息输入
def get_user_info():
    quarter = input("请输入季度 (例如: 1-4): ").strip()
    while quarter not in ['1', '2', '3', '4']:
        quarter = input("输入无效，请输入季度 (1-4): ").strip()

    name = input("请输入姓名: ").strip()
    number = input("请输入手机号: ").strip()
    
    # 保存到文件
    with open('userinfo.txt', 'w', encoding='utf-8') as f:
        f.write(f"quarter={quarter}\n")
        f.write(f"name={name}\n")
        f.write(f"number={number}\n")
    
    return {'quarter': quarter, 'name': name, 'number': number}

def load_user_info():
    try:
        user_info = {}
        with open('userinfo.txt', 'r', encoding='utf-8') as f:
            for line in f:
                key, value = line.strip().split('=')
                user_info[key] = value
        return user_info
    except FileNotFoundError:
        return get_user_info()

first_update = True

if __name__ == "__main__":
    # 创建参数解析器并添加说明
    parser = argparse.ArgumentParser(description="指定HTML文件夹路径和Word文件路径")
    parser.add_argument("html_folder", type=str, nargs='?', default=None, help="添加HTML文件夹路径")
    parser.add_argument("word_file", type=str, nargs='?', default=None, help="添加Word文件路径")

    args = parser.parse_args()
    html_folder_path, docx_file_path = args.html_folder, args.word_file
    
    # 如果html_folder_path或docx_file_path为空，则提示用户手动输入
    if not html_folder_path or not docx_file_path:
        user_input = input("请同时输入HTML文件夹路径和Word文件路径(空格分隔 在文件管理器中按住Shift键并同时点击右键,在弹出的菜单中选择\'复制为路径\'):\n")
        html_folder_path, docx_file_path = user_input.split()  # 使用空格分割用户输入内容
        # 去掉路径字符串中的双引号
        html_folder_path = html_folder_path.strip('"')
        docx_file_path = docx_file_path.strip('"')
    # html_folder_path,docx_file_path = 'html_folder','南京城市学院运维巡检记录-20240710.docx'# 硬编码添加路径
    
    # input(f"请确保已关闭 Word 文档，然后按 Enter 键继续...")

    # 复制文件到 save_folder
    new_path, new_html_folder_path = copy_files_to_save_folder(docx_file_path, html_folder_path, save_folder)
    
    # 遍历文件夹中的所有 HTML 文件
    for html_filename in os.listdir(new_html_folder_path):
        if html_filename.endswith('.html'):
            html_file_path = os.path.join(new_html_folder_path, html_filename)
            
            # 提取 IP 地址(文件名以 IP 开头)
            ip_address = html_filename.split('_')[0]
            
            # 加载和解析 HTML 文件
            html_content = load_html(html_file_path)
            check_items = parse_html_for_check_items(html_content)
            
            # 更新 DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
            if ip_address:
                logger.info(f"与 IP 地址 {ip_address} 匹配的表格内容:")
                update_docx_info(new_path, html_file_path, ip_address)
                matching_docx_tables(new_path, ip_address,check_items)
                logger.info(f"完成对 IP 地址 {ip_address} 的处理\n")
    # 打开修改后的文档
    # os.startfile(new_path)