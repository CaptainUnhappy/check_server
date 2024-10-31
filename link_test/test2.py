import re
import shutil
from bs4 import BeautifulSoup
from docx import Document
import os
import win32com.client as win32
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
from datetime import datetime
import time

# 获取当前时间并格式化为文件夹名
current_time = datetime.now().strftime("%Y%m%d_%H%M%S")

# 获取当前工作目录，并创建一个以时间命名的文件夹
current_dir = os.getcwd()
save_folder = os.path.join(current_dir, current_time)
os.makedirs(save_folder, exist_ok=True)

# 设置日志文件的路径
log_file = os.path.join(save_folder, f"{current_time}.log")

# 配置日志记录，确保日志输出到文件和控制台，且文件以UTF-8编码
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# 创建日志文件处理器，保存为UTF-8编码
file_handler = logging.FileHandler(log_file, encoding='utf-8')
file_handler.setLevel(logging.INFO)

# 创建控制台处理器
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)

# 设置日志格式
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# 将处理器添加到logger
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# 巡检内容列表
inspection_items = [
    "异常特权账户",
    "异常远程账户",
    "空密码账户",
    "爆破记录",
    "异常计划任务",
    "CPU使用率 < 80%",
    "内存使用率 < 80%",
    "TOP10进程信息",
    "僵尸进程",
    "服务器时间",
    "防火墙状态",
    "ESTABLISHED < 1000",
    "磁盘分区占用 < 80%",
]

manual_whitelist = []

# def get_manual_whitelist():
#     global manual_whitelist
#     # 允许用户手动输入白名单
#     imput_whitelist = input("请输入白名单用户和IP (以空格分隔)")
#     manual_whitelist += imput_whitelist.split()
#     logger.info(f'白名单用户和IP:{manual_whitelist}')
#     return manual_whitelist

def checked_all_correct(row,msg):
    row.cells[3].paragraphs[0].clear().add_run('✔')
    row.cells[4].paragraphs[0].clear()
    row.cells[5].paragraphs[0].clear()
    logger.info(f'    ✔ {msg}')

def checked_wrong_with_msg(row,msg):
    row.cells[3].paragraphs[0].clear().add_run('✘')
    row.cells[4].paragraphs[0].add_run(msg)
    logger.warning(f'✘ {msg}')


# 定义各个检查函数
def check_abnormal_privileged_accounts(row,check_items):
    item = check_items['用户安全审计'][0]
    # print(item)
    other = str(item).replace("特权用户列表：",'')
    msg = '异常特权账户 '
    if '' == other:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+other)


def check_abnormal_remote_accounts(row, check_items):
    item = check_items['用户安全审计'][2]
    other = str(item).replace("密码为空的用户列表：", '')

    # global manual_whitelist
    # if manual_whitelist:
    #     # 从列表中移除白名单
    #     for user in manual_whitelist:
    #         other = re.sub(rf'{user}\S*\b', '', other).strip()

    msg = '异常远程账户 '
    if '' == other:
        checked_all_correct(row, msg)
    else:
        checked_wrong_with_msg(row, msg + other)


def check_empty_password_accounts(row,check_items):
    item = check_items['用户安全审计'][2]
    # print(item)
    other = str(item).replace("密码为空的用户列表：",'')
    msg = '空密码账户 '
    if '' == other:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+other)


def check_brute_force_records(row, check_items):
    item = check_items['登陆失败记录']
    # print("原始记录:", item)

    cleaned_list = []
    recording = False

    for line in item:
        # 遇到 'IP              Failes' 开始记录
        if line == 'IP              Failes':
            recording = True
            continue
        # 遇到 '爆破主机root账号的可疑IP记录:' 停止记录第一部分
        if line == '爆破主机root账号的可疑IP记录:':
            recording = False
            continue
        # 如果正在记录，将内容添加到清理后的列表中
        if recording:
            cleaned_list.append(line)

    cleaned_item = cleaned_list

    # print("清理后的记录:", cleaned_item)
    global manual_whitelist
    if manual_whitelist:
        # 从列表中移除白名单
        cleaned_item = [line for line in cleaned_item if not any(whitelisted_ip in line for whitelisted_ip in manual_whitelist)]

    # print("清理后的记录（移除白名单后）:", cleaned_item)

    msg = '爆破记录 '
    if len(cleaned_item) == 0:
        checked_all_correct(row, msg)
    else:
        checked_wrong_with_msg(row, msg + "请手动检查")


def check_abnormal_scheduled_tasks(row,check_items):
    item = check_items['系统安全审计']
    # print(item)
    index = list(item).index("当前用户计划任务列表：")
    # print(start)
    task_list = []
    for i in range(index,len(item)):
        task_list.append(item[i])
    # print(task_list)
    msg = '异常计划任务 '
    if row.cells[3].paragraphs[0].text == '✔':
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+"请手动检查")


def check_cpu_usage(row,check_items):
    item = check_items['系统资源巡检区'][0]
    # print(item)
    max_usage = 80
    usage = item.replace('CPU使用率：','').replace('%','')
    if usage == '':
        usage = 0
    cpu_usage = float(usage)
    # print(cpu_usage)
    msg = 'CPU使用率 < 80% '
    if max_usage >= cpu_usage:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_memory_usage(row,check_items):
    item = check_items['配置信息']
    # print(item)
    index = next((i for i, s in enumerate(item) if '内存使用率' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    max_usage = 80
    usage = item.replace('内存使用率：','').replace('%','')
    ram_usage = float(usage)
    # print(ram_usage)
    msg = '内存使用率 < 80% '
    if max_usage >= ram_usage:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_top10_processes(row,check_items):
    logger.warning('手动检查 TOP10进程信息')
    return


def check_zombie_processes(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统当前僵尸进程数' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = item.replace("系统当前僵尸进程数：",'')
    # print(number)
    if number == '':
        number = 0
    number = int(number)
    msg = '僵尸进程 '
    if 0 == number:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_server_time(row,check_items):
    checked_all_correct(row,'服务器时间')


def check_firewall_status(row,check_items):
    checked_all_correct(row,'防火墙状态')


def check_established_connections(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统 established socket数量' in s), None)
    # print(index)
    item = item[index]
    # print(item)
    number = int(item.replace("系统 established socket数量: ",''))
    # print(number)
    max_connections = 1000
    msg = 'ESTABLISHED < 1000 '
    if max_connections > number:
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,item)


def check_disk_partition_usage(row,check_items):
    item = check_items['系统资源巡检区']
    # print(item)
    index = next((i for i, s in enumerate(item) if '系统磁盘分区存储使用情况' in s), None)
    # print(index)
    end = next((i for i, s in enumerate(item) if '系统当前进程数' in s), None)
    infos = []
    for i in range(index,end):
        infos.append(item[i])
    # print(infos)

    disk_list = []
    # 只保留包含 "%" 和 "/" 的行
    filtered_infos = [info for info in infos if '%' in info and '/' in info]
    for info in filtered_infos:
        parts = info.split()  # 将每一行按照空格分割
        if len(parts) > 5:  # 确保这一行有足够的字段
            try:
                if "%" in parts[5]:
                    use_percentage = parts[5].strip('%')  # 去掉 '%'
                elif "%" in parts[4]:
                    use_percentage = parts[4].strip('%')  # 去掉 '%'
                disk_list.append(float(use_percentage))  # 转换为整数
            except ValueError:
                logger.error(f'ValueError: in {disk_list}')

    # print(disk_list)
    max_percentage = 80
    msg = '磁盘分区占用 < 80% '
    if all(max_percentage > value for value in disk_list):
        checked_all_correct(row,msg)
    else:
        checked_wrong_with_msg(row,msg+"请手动检查")


# 加载 HTML 文件
def load_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# 解析 HTML 获取检查项和服务器 IP
def parse_html_for_check_items(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    check_items = {}
    
    # 遍历所有 <h2> 标签
    for h2 in soup.find_all('h2'):
        section_title = h2.get_text(strip=True)
        
        # 找到紧跟着的 <pre> 标签
        pre = h2.find_next('pre')
        if pre:
            pre_content = pre.get_text(strip=True)
            if pre_content:
                # 根据换行符分割内容
                lines = pre_content.split('\n')
                # 将每个条目添加到检查项字典中，标题为键
                check_items[section_title] = [line.strip() for line in lines if line.strip()]
    
    # print(check_items)
    return check_items

def insert_html_file_into_table_cell(doc_path, html_file_path, ip_address):
    # 获取当前脚本所在的文件夹路径
    current_directory = os.getcwd()

    # 组合文件的完整路径
    doc_path = os.path.join(current_directory, doc_path)
    html_file_path = os.path.join(current_directory, html_file_path)
    # print(f"{doc_path}  {html_file_path}  {ip_address}")

    # 打开 Word 应用
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False  # 设置为True可以看到Word界面

    try:
        # 打开现有的 Word 文档
        doc = word_app.Documents.Open(doc_path)

        # 遍历文档中的所有表格
        for table in doc.Tables:
            if "IP实例" in table.Cell(1, 1).Range.Text and "html报告" in table.Cell(1, 2).Range.Text:
                for row in range(2, len(table.Rows) + 1):  # 从第二行开始遍历，因为第一行是表头
                    # 获取左侧单元格的 IP 地址
                    ip_cell = table.Cell(row, 1).Range.Text.strip()

                    # 如果 IP 地址匹配
                    if ip_address in ip_cell:
                        logger.info(f"与 IP 地址 {ip_address} 匹配的表格行: 第 {row} 行")

                        # 获取对应行的右侧单元格（第二列）
                        right_cell = table.Cell(Row=row, Column=2)

                        # 检查HTML文件是否存在
                        if os.path.exists(html_file_path):
                            # 在右侧单元格中插入HTML文件
                            right_cell.Range.InlineShapes.AddOLEObject(
                                ClassType="htmlfile",
                                FileName=html_file_path,  # HTML文件路径
                                LinkToFile=False,
                                DisplayAsIcon=True,
                                # IconFileName="C:\\Windows\\System32\\shell32.dll",
                                IconFileName=r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
                                IconIndex=0,
                                IconLabel=os.path.basename(html_file_path)  # 文件图标下方显示的文件名
                            )
                            logger.info(f"文件已成功插入到与 IP 地址 {ip_address} 对应的右侧单元格中。")
                            return
                        else:
                            logger.warning(f"未找到IP地址 {ip_address} 对应的HTML文件: {html_file_path}")

        # # 保存并关闭文档
        doc.Save()
    except Exception as e:
        logger.error(f"打开或处理文档时出错: {e}")

    finally:
        # 确保关闭文档和 Word 应用
        if 'doc' in locals():  # 确保 doc 已成功打开
            doc.Close()
        word_app.Quit()

def check_matching_rows(matching_rows):
    for row in matching_rows:
            # print([cell.text for cell in row.cells])
            check_row = row.cells[2].text.strip()  # 巡检内容列
            # print(check_row)
            
            if check_row in inspection_items:
                
                if check_row == inspection_items[0]:
                    check_abnormal_privileged_accounts(row,check_items)
                    
                elif check_row == inspection_items[1]:
                    check_abnormal_remote_accounts(row,check_items)

                elif check_row == inspection_items[2]:
                    check_empty_password_accounts(row,check_items)
                    
                elif check_row == inspection_items[3]:
                    check_brute_force_records(row,check_items)

                elif check_row == inspection_items[4]:
                    check_abnormal_scheduled_tasks(row,check_items)

                elif check_row == inspection_items[5]:
                    check_cpu_usage(row,check_items)

                elif check_row == inspection_items[6]:
                    check_memory_usage(row,check_items)

                elif check_row == inspection_items[7]:
                    check_top10_processes(row,check_items)

                elif check_row == inspection_items[8]:
                    check_zombie_processes(row,check_items)

                elif check_row == inspection_items[9]:
                    check_server_time(row,check_items)

                elif check_row == inspection_items[10]:
                    check_firewall_status(row,check_items)

                elif check_row == inspection_items[11]:
                    check_established_connections(row,check_items)

                elif check_row == inspection_items[12]:
                    check_disk_partition_usage(row,check_items)


# DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
def matching_docx_tables(new_path, ip_address,check_items):
    doc = Document(new_path)
    matching_rows = []
    for table in doc.tables:
        for row in table.rows:
            # 只在表格的前两列中查找匹配的 IP 地址
            if len(row.cells) > 2 and ip_address in row.cells[1].text:
                matching_rows.append(row)
    # 匹配的行内容
    # print(matching_rows)

            # if len(row.cells) == 2:
            #     check_row = row.cells[0].text.strip()  # 巡检内容列
            #     print(check_row)
            #     insert_row = row.cells[1]
            #     # print(insert_row)
            #     insert_row.paragraphs.clear()
    
    if matching_rows:
        check_matching_rows(matching_rows)
    else:
        logger.warning('!!!no matching_rows!!!')

    # 保存修改后的文档
    try:
        doc.save(new_path)
    except PermissionError:
        logger.error("保存文档时出错，请确保 Word 文档未被占用，然后重新运行脚本。")
        exit(1)

def remove_html_tables(new_path, ip_address):
    doc = Document(new_path)
    for table in doc.tables:
        # 遍历表格中的每一行
        for row in table.rows:
            # 左侧单元格内容
            ip_cell = row.cells[0].text.strip()
            
            # 判断左侧是否包含IP地址
            if ip_address in ip_cell:
                # 删除右侧的html链接
                row.cells[1].text = ''  # 清空右侧单元格内容
                # 设置右侧单元格内容居中
                for paragraph in row.cells[1].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 保存修改后的文档
    try:
        doc.save(new_path)
    except PermissionError:
        logger.info("保存文档时出错，请确保 Word 文档未被占用，然后重新运行脚本。")
        exit(1)
        
# 复制 DOCX 和 HTML 文件到 save_folder
def copy_files_to_save_folder(docx_file_path, html_folder_path, save_folder):
    # 复制 DOCX 文件
    docx_dest_path = os.path.join(save_folder, os.path.basename(docx_file_path))
    shutil.copy(docx_file_path, docx_dest_path)
    # logger.info(f"已复制 DOCX 文件到 {docx_dest_path}")
    # 复制 DOCX 文件（用于检查）
    check_docx_dest_path = os.path.join(save_folder, 'check_' + os.path.basename(docx_file_path))
    shutil.copy(docx_file_path, check_docx_dest_path)
    logger.info(f"已备份并新建 DOCX 文件到 {check_docx_dest_path}")
    
    # 创建保存 HTML 文件的文件夹
    html_dest_folder = os.path.join(save_folder, 'html_files')
    os.makedirs(html_dest_folder, exist_ok=True)

    # 复制 HTML 文件
    for html_filename in os.listdir(html_folder_path):
        if html_filename.endswith('.html'):
            html_src_path = os.path.join(html_folder_path, html_filename)
            html_dest_path = os.path.join(html_dest_folder, html_filename)
            shutil.copy(html_src_path, html_dest_path)
            logger.info(f"已复制 HTML 文件到 {html_dest_path}")
    
    return check_docx_dest_path, html_dest_folder

if __name__ == "__main__":
    # 文件夹路径和 DOCX 文件路径
    html_folder_path = 'test_'
    docx_file_path = '南京城市学院运维巡检记录-20240710.docx'
    
    # 提示用户手动关闭 Word 程序
    input(f"请确保已关闭 Word 文档，然后按 Enter 键继续...")
    # get_manual_whitelist()

    # 复制文件到 save_folder
    new_path, new_html_folder_path = copy_files_to_save_folder(docx_file_path, html_folder_path, save_folder)
    
    # 遍历文件夹中的所有 HTML 文件
    for html_filename in os.listdir(new_html_folder_path):
        if html_filename.endswith('.html'):
            html_file_path = os.path.join(new_html_folder_path, html_filename)
            
            # 提取 IP 地址（假设文件名以 IP 开头）
            ip_address = html_filename.split('_')[0]
            
            # 加载和解析 HTML 文件
            html_content = load_html(html_file_path)
            check_items = parse_html_for_check_items(html_content)
            
            # 更新 DOCX 文件中与特定 IP 匹配的所有表格内容，仅保留包含对应 IP 的行
            if ip_address:
                logger.info(f"与 IP 地址 {ip_address} 匹配的表格内容:")
                matching_docx_tables(new_path, ip_address,check_items)
                remove_html_tables(new_path, ip_address)
                insert_html_file_into_table_cell(new_path, html_file_path, ip_address)
                logger.info(f"完成对 IP 地址 {ip_address} 的处理\n")
    # 打开修改后的文档
    # os.startfile(new_path)